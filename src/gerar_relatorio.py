import os
import pandas as pd
import numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def gerar_relatorio_abnt2(medias_produtivas, medias_improdutivas, col_equipe):
    """Gera relatório formatado em ABNT2 com análise das equipes."""
    
    REPORT_OUT = os.path.join('../result', 'relatorio_analise_equipes.docx')
    
    doc = Document()
    
    # Configuração da página (ABNT - A4)
    section = doc.sections[0]
    section.page_height = Inches(11.69)  # 297mm
    section.page_width = Inches(8.27)    # 210mm
    section.left_margin = Inches(1.18)   # 30mm
    section.right_margin = Inches(0.79)  # 20mm
    section.top_margin = Inches(1.18)    # 30mm
    section.bottom_margin = Inches(1.18) # 30mm
    
    # Título do relatório
    titulo = doc.add_heading('RELATÓRIO DE ANÁLISE DE DESEMPENHO DAS EQUIPES', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Data do relatório
    data_relatorio = doc.add_paragraph()
    data_relatorio.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    data_run = data_relatorio.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
    data_run.bold = True
    
    doc.add_paragraph()  # Espaço
    
    # Introdução
    intro = doc.add_paragraph()
    intro.add_run("1. INTRODUÇÃO\n").bold = True
    intro.add_run("Este relatório apresenta uma análise detalhada do desempenho das equipes operacionais, "
                 "com foco nos principais indicadores de produtividade e eficiência. A análise é dividida em "
                 "duas seções principais: registros produtivos e registros improdutivos, conforme classificação "
                 "do sistema. Para cada métrica, as equipes são classificadas da pior para a melhor performance, "
                 "com destaque para aquelas que apresentam desvios significativos em relação às metas estabelecidas.")
    
    doc.add_paragraph()  # Espaço
    
    # Metodologia
    metodologia = doc.add_paragraph()
    metodologia.add_run("2. METODOLOGIA\n").bold = True
    metodologia.add_run("As métricas foram calculadas com base nos registros de apontamento das equipes, "
                       "considerando os seguintes parâmetros:\n")
    metodologia.add_run("• TempExe_min: Tempo de execução (Liberada - No_Local) - Meta: 50min (produtivo) / 20min (improdutivo)\n").italic = True
    metodologia.add_run("• TempDesl_min: Tempo de deslocamento (No_Local - A_Caminho)\n").italic = True
    metodologia.add_run("• InterReg_min: Intervalo regulamentar (Fim_Intervalo - Início_Intervalo) - Meta: 60min\n").italic = True
    metodologia.add_run("• TempPrepEquipe_min: Tempo de preparação da equipe\n").italic = True
    metodologia.add_run("• Tempo de utilização: TempExe_min + TempDesl_min - Meta: 85% de 468min (397.8min)\n").italic = True
    metodologia.add_run("• Tempo ocioso: TempPrepEquipe_min + (60 - InterReg_min) ou TempPrepEquipe_min + 60 (se InterReg_min = 0)\n").italic = True
    
    doc.add_page_break()
    
    # Função auxiliar para adicionar tabela de ranking
    def adicionar_tabela_ranking(titulo, dados, coluna_valor, descricao="", meta=""):
        """Adiciona uma tabela de ranking ao documento."""
        doc.add_heading(titulo, level=3)
        
        if meta:
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Meta: {meta}\n").bold = True
        
        if descricao:
            desc_para = doc.add_paragraph(descricao)
        
        # Cria tabela com ranking
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Cabeçalho da tabela
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Posição'
        hdr_cells[1].text = 'Equipe'
        hdr_cells[2].text = 'Valor (min)'
        
        # Preenche tabela
        for idx, (equipe, valor) in enumerate(dados):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx + 1)
            row_cells[1].text = str(equipe)
            row_cells[2].text = f"{valor:.2f}"
        
        doc.add_paragraph()  # Espaço
    
    # Função para processar dados de uma seção
    def processar_secao(tipo, df_medias, num_secao):
        """Processa e adiciona uma seção completa ao relatório."""
        if df_medias is None or df_medias.empty:
            return
        
        doc.add_heading(f"{num_secao}. ANÁLISE DE REGISTROS {tipo}", level=1)
        
        # Filtra apenas as linhas de média geral
        df_geral = df_medias[df_medias[col_equipe].str.startswith('MédiaTodosDias', na=False)].copy()
        if df_geral.empty:
            doc.add_paragraph(f"Nenhum dado disponível para análise de registros {tipo}.")
            return
        
        # Remove prefixo "MédiaTodosDias" para ter o nome real da equipe
        df_geral['Equipe_Nome'] = df_geral[col_equipe].str.replace('MédiaTodosDias', '')
        
        # 1. TempExe_min
        if 'Media_TempExe_min' in df_geral.columns:
            df_temp_exe = df_geral[['Equipe_Nome', 'Media_TempExe_min']].copy()
            df_temp_exe = df_temp_exe.dropna()
            df_temp_exe = df_temp_exe.sort_values('Media_TempExe_min', ascending=False)  # Pior para melhor
            
            meta = "50 min" if tipo == 'PRODUTIVAS' else "20 min"
            desc = ("Esta métrica indica o tempo médio de execução das atividades. "
                   "Valores muito baixos podem indicar erro de apontamento nos momentos "
                   "'No_Local' e 'Liberada'.")
            
            dados_ranking = list(zip(df_temp_exe['Equipe_Nome'], df_temp_exe['Media_TempExe_min']))
            adicionar_tabela_ranking(f"{num_secao}.1 Tempo de Execução (TempExe_min)", 
                                    dados_ranking, desc, meta)
        
        # 2. TempDesl_min
        if 'Media_TempDesl_min' in df_geral.columns:
            df_temp_desl = df_geral[['Equipe_Nome', 'Media_TempDesl_min']].copy()
            df_temp_desl = df_temp_desl.dropna()
            df_temp_desl = df_temp_desl.sort_values('Media_TempDesl_min', ascending=False)  # Pior para melhor
            
            desc = ("Esta métrica indica o tempo médio de deslocamento. "
                   "Valores muito baixos podem indicar erro de apontamento nos momentos "
                   "'A_Caminho' e 'No_Local'.")
            
            dados_ranking = list(zip(df_temp_desl['Equipe_Nome'], df_temp_desl['Media_TempDesl_min']))
            adicionar_tabela_ranking(f"{num_secao}.2 Tempo de Deslocamento (TempDesl_min)", 
                                    dados_ranking, desc)
        
        # 3. Tempo de Utilização
        if 'Media_TempExe_min' in df_geral.columns and 'Media_TempDesl_min' in df_geral.columns:
            df_geral['Tempo_Utilizacao'] = df_geral['Media_TempExe_min'] + df_geral['Media_TempDesl_min']
            df_geral['Percentual_Utilizacao'] = (df_geral['Tempo_Utilizacao'] / 468) * 100
            
            df_utilizacao = df_geral[['Equipe_Nome', 'Tempo_Utilizacao', 'Percentual_Utilizacao']].copy()
            df_utilizacao = df_utilizacao.dropna()
            df_utilizacao = df_utilizacao.sort_values('Percentual_Utilizacao')  # Pior para melhor
            
            doc.add_heading(f"{num_secao}.3 Tempo de Utilização", level=3)
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Meta: 85% de 468min (397.8min)\n").bold = True
            
            desc_para = doc.add_paragraph("Tempo total de trabalho produtivo (execução + deslocamento). "
                                         "Valores abaixo de 85% indicam subutilização da jornada.")
            
            # Tabela
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Posição'
            hdr_cells[1].text = 'Equipe'
            hdr_cells[2].text = 'Tempo (min)'
            hdr_cells[3].text = 'Utilização (%)'
            
            for idx, row in df_utilizacao.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx + 1)
                row_cells[1].text = str(row['Equipe_Nome'])
                row_cells[2].text = f"{row['Tempo_Utilizacao']:.2f}"
                row_cells[3].text = f"{row['Percentual_Utilizacao']:.1f}%"
            
            doc.add_paragraph()
        
        # 4. InterReg_min
        if 'Media_InterReg_min' in df_geral.columns:
            df_interreg = df_geral[['Equipe_Nome', 'Media_InterReg_min']].copy()
            df_interreg = df_interreg.dropna()
            
            # Calcula desvio da meta (60min)
            df_interreg['Desvio_Meta'] = abs(df_interreg['Media_InterReg_min'] - 60)
            df_interreg = df_interreg.sort_values('Desvio_Meta', ascending=False)  # Pior para melhor
            
            doc.add_heading(f"{num_secao}.4 Intervalo Regulamentar (InterReg_min)", level=3)
            meta_para = doc.add_paragraph()
            meta_para.add_run("Meta: 60min (entre 4ª e 6ª hora)\n").bold = True
            
            desc_para = doc.add_paragraph("Intervalo para refeição. Desvios significativos podem indicar "
                                         "irregularidades no cumprimento da jornada.")
            
            # Tabela
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Posição'
            hdr_cells[1].text = 'Equipe'
            hdr_cells[2].text = 'Intervalo (min)'
            hdr_cells[3].text = 'Desvio da Meta'
            
            for idx, row in df_interreg.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx + 1)
                row_cells[1].text = str(row['Equipe_Nome'])
                row_cells[2].text = f"{row['Media_InterReg_min']:.2f}"
                row_cells[3].text = f"{row['Desvio_Meta']:.2f}"
            
            doc.add_paragraph()
        
        # 5. TempPrepEquipe_min
        if 'Media_TempPrepEquipe_min' in df_geral.columns:
            df_prep = df_geral[['Equipe_Nome', 'Media_TempPrepEquipe_min']].copy()
            df_prep = df_prep.dropna()
            df_prep = df_prep.sort_values('Media_TempPrepEquipe_min', ascending=False)  # Pior para melhor
            
            desc = ("Tempo de preparação da equipe. Valores elevados indicam possível ociosidade "
                   "ou ineficiência no processo de preparação para novas atividades.")
            
            dados_ranking = list(zip(df_prep['Equipe_Nome'], df_prep['Media_TempPrepEquipe_min']))
            adicionar_tabela_ranking(f"{num_secao}.5 Tempo de Preparação (TempPrepEquipe_min)", 
                                    dados_ranking, desc)
        
        # 6. Tempo Ocioso Total
        if 'Media_TempPrepEquipe_min' in df_geral.columns and 'Media_InterReg_min' in df_geral.columns:
            # Calcula tempo ocioso: TempPrepEquipe_min + (60 - InterReg_min) ou TempPrepEquipe_min + 60 se InterReg_min = 0
            df_geral['Tempo_Ocioso'] = np.where(
                df_geral['Media_InterReg_min'] == 0,
                df_geral['Media_TempPrepEquipe_min'] + 60,
                df_geral['Media_TempPrepEquipe_min'] + (60 - df_geral['Media_InterReg_min'])
            )
            
            df_ocioso = df_geral[['Equipe_Nome', 'Tempo_Ocioso']].copy()
            df_ocioso = df_ocioso.dropna()
            df_ocioso = df_ocioso.sort_values('Tempo_Ocioso', ascending=False)  # Pior para melhor
            
            doc.add_heading(f"{num_secao}.6 Tempo Ocioso Total", level=3)
            desc_para = doc.add_paragraph("Soma do tempo de preparação com o tempo não utilizado do intervalo. "
                                         "Valores elevados indicam ociosidade operacional significativa.")
            
            # Tabela
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Posição'
            hdr_cells[1].text = 'Equipe'
            hdr_cells[2].text = 'Tempo Ocioso (min)'
            
            for idx, row in df_ocioso.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx + 1)
                row_cells[1].text = str(row['Equipe_Nome'])
                row_cells[2].text = f"{row['Tempo_Ocioso']:.2f}"
            
            doc.add_paragraph()
        
        doc.add_page_break()
    
    # Processa seções
    secao_num = 3
    if medias_produtivas is not None and not medias_produtivas.empty:
        processar_secao('PRODUTIVAS', medias_produtivas, secao_num)
        secao_num += 1
    
    if medias_improdutivas is not None and not medias_improdutivas.empty:
        processar_secao('IMPRODUTIVAS', medias_improdutivas, secao_num)
        secao_num += 1
    
    # Conclusão
    doc.add_heading(f"{secao_num}. CONCLUSÕES E RECOMENDAÇÕES", level=1)
    conclusao = doc.add_paragraph()
    conclusao.add_run("Com base na análise realizada, observa-se que:\n\n").bold = True
    conclusao.add_run("1. As equipes com pior desempenho nas métricas de tempo devem receber atenção especial;\n")
    conclusao.add_run("2. Valores muito abaixo do padrão em TempExe_min e TempDesl_min sugerem necessidade de "
                     "treinamento sobre apontamento correto;\n")
    conclusao.add_run("3. Tempos ociosos elevados indicam oportunidades de melhoria na gestão operacional;\n")
    conclusao.add_run("4. Desvios significativos no intervalo regulamentar requerem verificação do cumprimento "
                     "da jornada de trabalho;\n")
    conclusao.add_run("5. Recomenda-se acompanhamento periódico destes indicadores para melhoria contínua.\n\n")
    conclusao.add_run("Este relatório deve ser utilizado como base para planos de ação corretivos e preventivos.")
    
    # Salva o documento
    doc.save(REPORT_OUT)
    print(f"Relatório ABNT2 gerado: {REPORT_OUT}")
    
    return doc